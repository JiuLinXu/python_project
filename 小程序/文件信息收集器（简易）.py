'''
1、图形化界面；
2、选择一个文件夹，读取这个文件夹及其所有子文件夹中的所有文件； 
3、获取所有文件的文件路径、文件主名、文件扩展名、文件大小（KB）、最后修改时间等信息； 
4、如果是word文件、txt文件或者pdf文件则还需读取文件的字数、段落数等；
5、如果是Excel文件则需要读取表格的sheet数，和表格的行列数；
6、文件收集过程采用进度条；
7、无法读取的文件（如：快捷方式）则跳过；
8、最后将结果生成一个Excel表。
'''
import os
import pandas as pd
from tkinter import Tk, Button, filedialog, messagebox, Text, Scrollbar
from tkinter.ttk import Progressbar
from datetime import datetime
from docx import Document
# import PyPDF2
import pdfplumber
import win32com.client
def get_file_info(file_path):
    try:
        file_info = {
            "路径": file_path,
            "主名": os.path.basename(file_path),
            "扩展名": os.path.splitext(file_path)[1],
            "大小 (KB)": round(os.path.getsize(file_path) / 1024,1),  # Convert bytes to KB  
            "最后修改时间": datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')  
        }

        # 处理 Word 和文本文件
        if file_info["扩展名"] in ['.docx', '.txt']:  
            doc = Document(file_path)
            word_count = 0
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                # 计算每一段的字数，使用strip()以解去空格，split()按空格分词  
                if paragraph.text.strip():  # 如果段落不为空  
                    word_count += len(paragraph.text.split())  
                    paragraph_count += 1
            # 统计列表内容 (如项目符号、编号列表)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        word_count += len(cell.text.split())  
            file_info["字数"] = word_count  # 设置字数
            file_info["段落数"] = paragraph_count  # 设置段落数 
        # 处理 .doc 文件  
        elif file_info["扩展名"] == '.doc':  
            word = win32com.client.Dispatch("Word.Application")  
            doc = word.Documents.Open(file_path)  
            file_info["字数"] = doc.Words.Count  # 获取字数  
            file_info["段落数"] = doc.Paragraphs.Count  # 获取段落数  
            doc.Close(False)  # 关闭文档，不保存更改  
            word.Quit()  # 退出Word应用  
        # 处理 PDF 文件  
        elif file_info["扩展名"] == '.pdf':  
            try:
                with pdfplumber.open(file_path) as pdf:  
                    text = ""  
                    for page in pdf.pages:  
                        text += page.extract_text() + "\n"  # 提取文本并添加换行  
                    file_info["字数"] = len(text.split())  
                    file_info["段落数"] = text.count('\n')  # 段落数以换行符计算  
            except Exception as e:
                print(file_path)
                print(f"读取PDF时出错: {e}")    

        # 处理 Excel 文件  
        elif file_info["扩展名"] in ['.xlsx', '.xls']:  
            xls = pd.ExcelFile(file_path)  
            file_info["sheet数"] = len(xls.sheet_names)  
            file_info["行数"] = sum(pd.read_excel(xls, sheet).shape[0] for sheet in xls.sheet_names)  
            file_info["列数"] = sum(pd.read_excel(xls, sheet).shape[1] for sheet in xls.sheet_names)  
        return file_info  
    except Exception as e:  
        return None  
def collect_files(folder_path, progress_bar, text_box):  
    file_data = []  
    files_to_process = []  
    # 收集所有文件路径
    for root, dirs, files in os.walk(folder_path):  
        files_to_process.extend([os.path.join(root, file) for file in files])  
    progress_bar['maximum'] = len(files_to_process)  # 设置进度条最大值  
    # 遍历文件并更新进度条和文本框  
    for file_path in files_to_process:  
        info = get_file_info(file_path)  
        if info:  
            file_data.append(info)  
            # 在文本框中显示文件路径和名称  
            text_box.insert('end', f"路径: {info['路径']}\n")  
            text_box.insert('end', f"主名: {info['主名']}\n")  
            text_box.insert('end', f"扩展名: {info['扩展名']}\n")  
            text_box.insert('end', f"大小 (KB): {info['大小 (KB)']}\n")  
            text_box.insert('end', f"最后修改时间: {info['最后修改时间']}\n")  
            text_box.insert('end', "-" * 40 + "\n")  # 分隔线  
        progress_bar['value'] += 1  # 更新进度条  
        progress_bar.update()  # 强制更新进度条显示  
    return file_data  
def save_to_excel(data, output_file):  
    # 创建一个 Excel writer 对象  
    with pd.ExcelWriter(output_file, engine='xlsxwriter') as writer:  
        # 将所有数据写入“总表”  
        df = pd.DataFrame(data)  
        df.to_excel(writer, sheet_name='总表', index=False)  # 写入“总表”  
        # 根据文件扩展名分类数据  
        categorized_data = {}  
        for item in data:  
            ext = item["扩展名"]  
            if ext not in categorized_data:  
                categorized_data[ext] = []  
            item["打开文件"] = f'=HYPERLINK("{item["路径"]}", "打开文件")'  # 添加超链接  
            categorized_data[ext].append(item)  
        # 为每种扩展名创建单独的工作表  
        for ext, items in categorized_data.items():  
            df_ext = pd.DataFrame(items)  
            # 创建工作表名，去掉“.”前缀  
            sheet_name = ext[1:] if ext.startswith('.') else '无扩展名'  
            df_ext.to_excel(writer, sheet_name=sheet_name, index=False)  
        messagebox.showinfo("完成", f"文件信息已保存到 {output_file}")  
def select_folder(progress_bar, text_box):  
    folder_path = filedialog.askdirectory()  
    if folder_path:  
        progress_bar['value'] = 0  # 重置进度条  
        text_box.delete(1.0, 'end')  # 清空文本框  
        file_data = collect_files(folder_path, progress_bar, text_box)  
        if file_data:  
            output_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])  
            if output_file:  
                save_to_excel(file_data, output_file)  
        else:  
            messagebox.showinfo("信息", "未找到可读取的文件。")  
def main():  
    root = Tk()  
    root.title("文件信息收集器")  
    root.geometry("600x400")  
    # 创建选择文件夹按钮  
    select_button = Button(root, text="选择文件夹", command=lambda: select_folder(progress_bar, text_box))  
    select_button.pack(pady=10)  
    # 创建进度条  
    progress_bar = Progressbar(root, orient="horizontal", length=500, mode="determinate")  
    progress_bar.pack(pady=10)  
    # 创建文本框以显示文件信息  
    text_box = Text(root, wrap='word', width=70, height=15)  
    text_box.pack(pady=10)  
    # 添加滚动条  
    scroll_bar = Scrollbar(root, command=text_box.yview)  
    scroll_bar.pack(side='right', fill='y')  
    text_box.config(yscrollcommand=scroll_bar.set)  
    root.mainloop()  
if __name__ == "__main__":  
    main()