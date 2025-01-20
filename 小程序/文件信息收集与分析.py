# coding=UTF-8
#qwererr
import os,time,sys,shutil
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
import re
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

sys.path.append("/python-project")
from service.My_tools import TextAnalysis,writeExcel,Opration_Word,Opration_PDF,Opration_Similarity,ReadExcel_pandas
class CollectionAndAnalysisOfFileInformation:
    def __init__(self, root):
        
        # projectName="根据文本内容，获得文本的相似度" 
        # projectName="对文本内容进行分析，将分析结果保存到Excel表" #
        # projectName="根据Excel表的内容将文件更改文件名后，拷贝到新的文件夹，同时生成总表" #规范文件名，以便评选
        # projectName="批量发送电子邮件" #如给评委分发电子邮件
        # projectName="按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中" #评委评选完成按照获奖级别等奖文件进行分类
        # projectName="比较两个文本或文档_比较两个word文档，并获取修订"

        self.root = root
        self.root.title("文件信息收集与分析智能体")
        self.root.geometry("900x600")
        # 创建 ttk.Style 对象
        style = ttk.Style()
        # 配置字体样式
        style.configure("TCombobox", font=("黑体", 14))  # 设置字体为 Arial，字号为 14
        # 创建下拉选项框
        self.options = [
            "收集文件信息，生成Excel表格",
            "根据表格已有字段（如：文件路径、文件名）补充完善信息",
            "从申报表中采集信息",
            "根据指定字段，筛选重复记录",
            "对文本内容进行分析，将分析结果保存到Excel表",
            "根据Excel表的内容将文件更改文件名后，拷贝到新的文件夹，同时生成总表",
            "分析文本相似度",
            "批量发送电子邮件",
            "按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中",
            "比较两个文本或文档_比较两个word文档，并获取修订"
        ]
        self.combobox = ttk.Combobox(root, values=self.options, state="readonly", width=100,style="TCombobox")
        self.combobox.pack(pady=10)
        self.combobox.bind("<<ComboboxSelected>>", self.on_combobox_select)
        
        # 进度条
        self.progress = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
        self.progress.pack(pady=20)

        # 日志显示
        self.log_text = tk.Text(root, height=10, width=50)
        self.log_text.pack(pady=10)
    
    def on_combobox_select(self, event):
        selected_option = self.combobox.get()
        if selected_option == "收集文件信息，生成Excel表格":
            self.collect_files()
        elif selected_option == "分析文本相似度":
            self.analyze_similarity()
        elif selected_option == "批量发送电子邮件":
            self.send_emails()
    def collect_files(self):
        folder_path = filedialog.askdirectory(title="选择文件夹")
        if folder_path:
            self.log("开始收集文件信息...")
            self.progress["value"] = 0
            self.root.update_idletasks()
            # 调用文件信息收集函数
            file_info = self.get_file_info(folder_path)
            save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if save_path:
                file_info.to_excel(save_path, index=False)
                self.log(f"文件信息已保存到: {save_path}")
                self.progress["value"] = 100

    def get_file_info(self, folder_path):
        file_info = []
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_name, file_extension = os.path.splitext(file)
                file_info.append({
                    "文件路径": root,
                    "文件名": file_name,
                    "扩展名": file_extension,
                    "打开文件": f'=HYPERLINK("{file_path}", "打开文件")'
                })
        return pd.DataFrame(file_info)

    def analyze_similarity(self):
        file_path = filedialog.askopenfilename(title="选择Excel文件", filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            self.log("开始分析文本相似度...")
            self.progress["value"] = 0
            self.root.update_idletasks()

            # 调用文本相似度分析函数
            df = pd.read_excel(file_path)
            similarity_results = self.calculate_similarity(df)
            save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if save_path:
                similarity_results.to_excel(save_path, index=False)
                self.log(f"相似度分析结果已保存到: {save_path}")
                self.progress["value"] = 100

    def calculate_similarity(self, df):
        # 这里可以调用之前的相似度计算函数
        df["相似度"] = "待实现"
        return df

    def send_emails(self):
        self.log("开始批量发送电子邮件...")
        self.progress["value"] = 0
        self.root.update_idletasks()

        # 调用邮件发送函数
        sender = ("69301020@qq.com", "dceyjpkswzhhbhbf")
        recipients = [("look.east@163.com", "张向东"), ("kmswgy33@126.com", "张建欣")]
        subject = "“中华魂”（毛泽东伟大精神品格）征文评选"
        content = "尊敬的评委老师，您好！请查收附件。"

        for recipient in recipients:
            try:
                server = zmail.server(*sender)
                mail_msg = {
                    "subject": subject,
                    "content_text": content,
                    "attachments": ["path/to/attachment.rar"]
                }
                server.send_mail(recipient[0], mail_msg)
                self.log(f"邮件已发送至: {recipient[1]}")
                self.progress["value"] += 100 / len(recipients)
                self.root.update_idletasks()
            except Exception as e:
                self.log(f"发送邮件失败: {str(e)}")
    def log(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
class Main():
    def main():
        '''
        工作步骤：
        前期下发文件时，最好要求征文为Word文档、报名表为Excel文档，上报的文件名称统一为“学校名称_作品名称_学生姓名_指导教师_联系电话.后缀”；
        1、收取电子邮件附件，解压，按照后缀分类（视频、Word文档、Excel文档。如果前期报名表为Word则需将Word文档又分为作品和报名表）；
        2、采用"收集文件信息，可生成Excel表格"，分别针对保存视频、Word文档的文件夹进行文件信息收集，生成“视频统计表”和“征文统计表”；
        3、将所有表名进行文档合并，生成“报名表汇总表”；
        4、根据表格中的文件信息结合“报名表汇总表”完善“视频统计表”和“征文统计表”，例如：生成“作品编号”、“学校”、“姓名”、“题目”、“指导教师姓名”、“联系电话”等信息）；
        5、按照完善好的“视频统计表”和“征文统计表”，对所有文件名进行重新命名。（"按照Excel表的内容将文件拷贝到新的文件夹，并更改文件名"）；
        6、制作“评委评分统计总表”，该总表在完善好的“视频统计表”和“征文统计表”上进行删减生成，尽量提供一些信息给评委（如：“字数”、“段落”、“文档相似度”、“最相似的文章名称”等）；
        7、按照评委分工情况，将“评委评分统计总表”进行分解，生成每一位评委使用的“评委评分统计表”，同时将需要评审的作品文件拷贝到一起，同时将“评分说明”放入，打包；
        8、使用"批量发送电子邮件"给每一位评委发电子邮件；
        9、收集每一位评委填写好的“评委评分统计表”，对接到“评委评分统计总表”中；
        10、按照“评委评分统计总表”，用"按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中" ，将评委评选完成按照获奖级别等文件进行分类；
        11、按照“评委评分统计总表”生成并打印奖状。
        '''
        projectName="收集文件信息，生成Excel表格" #收集文件信息，包含：文件名、扩展名……第一行文字……、总字数（参考）……
        # projectName="根据表格已有字段（如：文件路径、文件名）补充完善信息" #收集电子邮件发来并下载后的所有文件信息
        # projectName="从申报表中采集信息" 
        # projectName="根据指定字段，筛选重复记录"
        # projectName="根据文本内容，获得文本的相似度" 
        # projectName="对文本内容进行分析，将分析结果保存到Excel表" #
        # projectName="根据Excel表的内容将文件更改文件名后，拷贝到新的文件夹，同时生成总表" #规范文件名，以便评选
        # projectName="批量发送电子邮件" #如给评委分发电子邮件
        # projectName="按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中" #评委评选完成按照获奖级别等奖文件进行分类
        # projectName="比较两个文本或文档_比较两个word文档，并获取修订"
        if projectName=="收集文件信息，生成Excel表格":
            # path_name=r"D:\python-project\关工委\20240912从小学党史活动\下载的附件\附件汇总（有效）"
            # path_name=r'E:\周工作小结\20241121五华区数智竞赛'
            # path_name=r'E:\周工作小结\20241030课程实施办法和方案'
            # path_name=r'D:\GPT浏览器下载\2025-01-07—2025-01-10（609918474@qq.com）'
            path_name=r'C:\Users\XJL\Desktop\教学设计'
            save_pathfile=r"E:\1.xlsx"
            #获取参数
            para=Main.cottent_para_set() #统计与分析参数设定
            CollectFileInformation.main( path_name,save_pathfile,para)
        elif projectName=="根据表格已有字段（如：文件路径、文件名）补充完善信息":
            path_name=r"D:\python-project\关工委\20240912从小学党史活动"
            excel_file='汇总表.xlsx'
            sheet_name='文件信息采集表'
            out_pathfile=r"E:\2.xlsx"
            FieldsInformation.main(path_name,excel_file,sheet_name,out_pathfile)
        elif projectName=="从申报表中采集信息": #完善汇总表信息
            excel_path_file=r'D:\python-project\关工委\20240912从小学党史活动\汇总表.xlsx'
            file2=r'D:\python-project\关工委\20240912从小学党史活动\1.xlsx'
            df1=pd.read_excel(excel_path_file,sheet_name='汇总表去重', header=0)
            df1 = df1[(df1['县区'] == '明致')] #& (pd.isna(df1['学校']))]
            df2=pd.read_excel(file2,sheet_name='Sheet1', header=0)
            save_file_path=r'E:\11.xlsx'
            GetInformationFromTable.main(df1,df2,save_file_path)
        elif projectName=="根据指定字段，筛选重复记录":
            excel_path_file=r'E:\1.xlsx'
            sheet_name='Sheet1'
            save_duplicates=r'E:\22.xlsx' #重复的所有记录
            save_unique=r'E:\33.xlsx' #清除重复行（仅保留第一行）后的所有数据
            filter_field=['文件名', '扩展名']
            FilterDuplicateRecords.main(excel_path_file,sheet_name,filter_field,save_duplicates,save_unique)
        elif projectName=="根据文本内容，获得文本的相似度":
            excel_path_file=r'E:\1.xlsx' #此Excel中需要有：文件路径、文件名、扩展名、打开文件
            sheet_name='Sheet1'
            text_ID="序号" #每一行的唯一ID，如：编号
            file_info=['文件路径','文件名','扩展名'] #
            save_file=r'E:\54.xlsx'
            GetTextSimilarity.main(excel_path_file,sheet_name,text_ID,file_info,save_file)
        elif projectName=="对文本内容进行分析，将分析结果保存到Excel表":
            excel_path_file=r'E:\1.xlsx' #此Excel中需要有：文件路径、文件名、扩展名、打开文件
            sheet_name='Sheet1'
            file_info=['文件路径','文件名','扩展名'] #
            save_file=r'E:\54.xlsx'
            x=TextAnalysis.main(excel_path_file,sheet_name,file_info)
        elif projectName=="根据Excel表的内容将文件更改文件名后，拷贝到新的文件夹，同时生成总表":
            excel_path_file=r"D:\python-project\关工委\20240912从小学党史活动\汇总表.xlsx"
            newPath=r"E:\1\征文"
            newNameCombo=['_',("编号","组别","征文标题")] #[分隔符，（文件名选取的字段元组）]
            sheet_name='汇总表去重'
            GenerateNewFilename.main(excel_path_file,newPath,newNameCombo,sheet_name)
        elif projectName=="批量发送电子邮件":
            SendEmail.SendEmail1()
        elif projectName=="按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中":
            excel_path_file=r'E:\20230109中华魂演讲、征文作品\20240117征文（送审）\征文（送审）\分解并发送电子邮件给评委\6征文（送审）\6“中华魂”（毛泽东伟大精神品格）征文初评表.xlsx'
            fileNameCol='征文序号' #文件名包含字符的列名
            classifyCol=['一等奖','二等奖','三等奖','优秀奖']#分类到文件夹的列
            filePath=r'E:\20230109中华魂演讲、征文作品\20240117征文（送审）\征文（送审）\分解并发送电子邮件给评委\6征文（送审）' #被分类的文件所在路径
            DirClassify.dirClassify1(excel_path_file,fileNameCol,classifyCol,filePath)
            pass
        elif projectName=="比较两个文本或文档_比较两个word文档，并获取修订":
            pathDoc1=r"E:\20230109中华魂演讲、征文作品\20240119新增的征文、视频、报名表\征文\伟人精神代代传.docx"
            pathDoc2=r"E:\20230109中华魂演讲、征文作品\征文\6.昆明市西山区书林二小'中华魂'活动  征文 参赛作品及报名表\昆明市西山区书林二小'中华魂'活动  征文 参赛作品及报名表\书二福海“中华魂”征文 报送材料\书二福海 征文作品\三（2）赵堇言《伟人精神代代传》.docx"
            Opration_Similarity.comparedDocx1(pathDoc1,pathDoc2)
        elif projectName=="比较word文档的相似度":
            pathDoc1=r"E:\20230109中华魂演讲、征文作品\20240119新增的征文、视频、报名表\征文\伟人精神代代传.docx"
            pathDoc2=r"E:\20230109中华魂演讲、征文作品\征文\6.昆明市西山区书林二小'中华魂'活动  征文 参赛作品及报名表\昆明市西山区书林二小'中华魂'活动  征文 参赛作品及报名表\书二福海“中华魂”征文 报送材料\书二福海 征文作品\三（2）赵堇言《伟人精神代代传》.docx"
            if os.path.splitext(pathDoc1)[1].lower()=='.doc':
                x2=Opration_Word.doc2docx(pathDoc1,os.getenv('TEMP'))
                wordTxt1=Opration_Word.readWord(os.path.join(os.getenv('TEMP'),x2))
                os.remove(os.path.join(os.getenv('TEMP'),x2))
            else:
                wordTxt1=Opration_Word.readWord(os.path.join(pathDoc1))
            if os.path.splitext(pathDoc2)[1].lower()=='.doc':
                x2=Opration_Word.doc2docx(pathDoc2,os.getenv('TEMP'))
                wordTxt2=Opration_Word.readWord(os.path.join(os.getenv('TEMP'),x2))
                os.remove(os.path.join(os.getenv('TEMP'),x2))
            else:
                wordTxt2=Opration_Word.readWord(os.path.join(pathDoc2))
            x1=Opration_Similarity.similarityText1(''.join(wordTxt1['bodyText']),''.join(wordTxt2['bodyText']))
            print(x1)
    #统计与分析参数设定
    def cottent_para_set():
        #文件属性
        file_attributes=['大小（KB）',
                         '最后修改时间',
                        #  '创建时间',
                        #  '最后访问时间',
                        #  '文件权限',
                         ]
        #常规字段统计
        narrow_filed=[
            '总字数（参考）',
            '有文字的段落数（参考）',
            ]
        #文本分析的相关指标
        role_system='''你是“AI（人工智能）+教学”教学竞赛的资深评委，你的任务是为用户提供专业、准确、有见地的建议并给出评分结果。'''
        text_analysis={
            'SnowNLP':{
                "分词结果":False,
                "情感得分":True,
                "转换成拼音":False,  #将文本转换成拼音表示。注意，这个功能可能需要额外的安装或配置，因为 SnowNLP 的标准安装可能不包括拼音转换功能。
                "关键词":True,  #将文本转换成拼音表示。注意，这个功能可能需要额外的安装或配置，因为 SnowNLP 的标准安装可能不包括拼音转换功能。
                "摘要":False,
            },
            'Jieba':{},
            'HanLP':{},
            'GAI':{
                'DeepSeek':role_system,
                # 'KiMi':role_system,
                # 'XunFei':role_system,
                # 'ChatGPT':role_system,
                # 'zhipuai':role_system,
                # 'volcano':role_system,
            },
        }
        # 采用正则表达是的特殊字符及其正则表达式，无需获取可为空
        re_spec_dict = {
            # '学校': r"(?<!从)(?<=\S)[\u4e00-\u9fa5]+(小学|中学|学校)",  
            '姓名': r"(学生姓名：|学生：|\s)([\u4e00-\u9fa5]{2,4})(?=\s|$)",  
            # '班级': r"(初|高)[一二三四五六七八九十]+班",  
            # '指导教师': r"指导教师?：?([\u4e00-\u9fa5]{2,4})",  
            '联系电话': r"\b1[3-9]\d{9}\b"  
        }
        # 指定字符前后的字符
        specified_before_after={
            '依据':[
                'before',   #before：此字符之前；after：此字符之后；in：其中有此字符
                True,       # （上一参数为in时无效）True：截取字符包含此字符；False：截取字符不包含此字符
                -1,         # >0：指定数量的字符；-1：截取整个段落
                'endone',   #endone：整篇文章最后一次出现；firstone：整篇文章第一次出现；all：整篇文章所有
            ], 
            '中心':['after',True,-1,'firstone'],
        }
        #获取指定段落内容，key为段落的编号
        paragraph_text={
            # 1:'第一段文字',
            # 2:'第二段文字',
            # 3:'第三段文字'
            }
        return [file_attributes,narrow_filed,text_analysis,re_spec_dict,specified_before_after,paragraph_text]
    

#根据表格已有字段（如：文件路径、文件名）补充完善信息
class FieldsInformation():
    def main(path,excel_file,sheet_name,out_pathfile):
        import pandas as pd
        kunming_districts = ["五华","盘龙","官渡","西山","东川","呈贡","晋宁","安宁","富民","宜良","嵩明","石林","禄劝","寻甸"]
        file_path = os.path.join(path, excel_file)  
        df=pd.read_excel(file_path,sheet_name=sheet_name, header=0, index_col=0)
        df['县区'] = df['县区'].astype(str)
        df['县区'] = None
        #正则表达式
        # school_pattern =  r'(?!从小学).*?(?:小学|中学|学校)' # 匹配学校名称
        school_pattern =  r'(?!从小).*?(?:小)' # 匹配学校名称
        surnames = ["李", "王", "张", "刘", "陈", "杨", "赵", "黄", "周", "吴", "徐", "孙", "胡", "朱", "高", "林", "何", "郭", "马", "罗", "梁", "宋", "郑", "谢", "韩", "唐", "冯", "于", "董", "萧", "程", "曹", "袁", "邓", "许", "傅", "沈", "曾", "彭", "吕", "苏", "卢", "蒋", "蔡", "贾", "丁", "魏", "薛", "叶", "阎", "余", "潘", "杜", "戴", "夏", "钟", "汪", "田", "任", "姜", "范", "方", "石", "姚", "谭", "廖", "邹", "熊", "金", "陆", "郝", "孔", "白", "崔", "康", "毛", "邱", "秦", "江", "史", "顾", "侯", "邵", "孟", "龙", "万", "段", "雷", "钱", "汤", "尹", "黎", "易", "常", "武", "乔", "贺", "赖", "龚", "文", "庞", "樊", "兰", "殷", "施", "陶", "洪", "翟", "安", "颜", "倪", "严", "牛", "温", "芦", "季", "俞", "章", "鲁", "葛", "伍", "韦", "申", "尤", "毕", "聂", "丛", "焦", "向", "柳", "邢", "路", "岳", "齐", "沿", "梅", "莫", "庄", "辛", "管", "祝", "左", "涂", "谷", "祁", "时", "舒", "耿", "牟", "卜", "路", "詹", "关", "苗", "凌", "费", "纪", "靳", "盛", "童", "欧", "毕", "郝", "邬", "安", "常", "乐", "于", "时", "傅", "皮", "卞", "齐", "康", "伍", "余", "元", "卜", "顾", "孟", "平", "黄", "和", "穆", "萧", "尹"]
        name_pattern = r'(?:' + '|'.join(map(re.escape, surnames)) + r')[\u4e00-\u9fa5]+'
        # name_pattern =  r'[\u4e00-\u9fa5]{2,4}'  # 匹配条件下的2到4个汉字姓名
        # name_pattern = r"学生姓名?：|学生?：?([\u4e00-\u9fa5]{2,4})"  # 匹配指导教师姓名
        # class_pattern = r'^[一二三四五六七八九初高cgCG][\s\S]*班$'
        # r'([一二三四五六七八九初高cgCG].*?班)'
        class_pattern = r'([一二三四五六七八九初高123456789cgCG].*?班)'# 匹配班级（如“初一3班”）
        teacher_pattern = r"指导教师?：?([\u4e00-\u9fa5]{2,4})"  # 匹配指导教师姓名
        phone_pattern = r"\b1[3-9]\d{9}\b"   # 匹配手机号码
        school_series = df['学校'].dropna().drop_duplicates() #学校列去除Nan、去重
        for row in df.itertuples(index=True, name='Pandas'):
            str1=row.文件路径.replace(r'D:\python-project\关工委\20240912从小学党史活动\下载的附件\附件汇总（有效）','')
            str2=row.文件名
            school_info = str1 + str2 + str(row.学校)
            val_list=[]
            #按照正则表达式获取
            # val_list = val_list+re.findall(teacher_pattern, str1)
            # val_list = val_list+re.findall(phone_pattern, str1)
            # val_list = val_list+re.findall(teacher_pattern, str2)
            # val_list = val_list+re.findall(phone_pattern, str2)

            #填写县区
            if pd.isna(row.县区):
                for district in kunming_districts:
                    if district in school_info:
                        df.at[row.Index, '县区'] = district
                        break
            if '工作情况' in row.文件名:
                df.at[row.Index, '组别'] = '工作情况'
                df.at[row.Index, '学校'] = ''
                df.at[row.Index, '姓名'] = ''
                df.at[row.Index, '班级'] = ''
            elif '申报表' in row.文件名 or '报名表' in row.文件名 or '汇总表' in row.文件名:
                df.at[row.Index, '组别'] = '申报表'
                df.at[row.Index, '学校'] = ''
                df.at[row.Index, '姓名'] = ''
                df.at[row.Index, '班级'] = ''
            else:
                #填写学校
                # if pd.isna(row.学校):
                #     #已有学校名获取
                #     for school_name in school_series:
                #         if school_name in row.文件路径 or school_name in row.文件名:
                #             df.at[row.Index, '学校']=school_name
                #     #正则表达式获取
                #     if pd.isna(row.学校):
                #         val_list = val_list+re.findall(school_pattern, str1)
                #         val_list = val_list+re.findall(school_pattern, str2)
                #         if len(val_list)>0:
                #             #删除列表中已经被其他元素包含了的元素
                #             val_list=FieldsInformation.clear_val_list(val_list)
                #             # 使用 filter 函数删除符合条件的元素  使用逻辑取反保留不符合条件的元素
                #             condition = lambda x: '从小' in x or x =='小' or x =='小学' or x =='（小学' or '征文' in x or '党史”征文' in x or '党史“征文' in x or '.' in x
                #             val_list = list(filter(lambda x: not condition(x), val_list))
                #             # 填写到df中
                #             df.at[row.Index, '学校'] = '|'.join(val_list)
                # 填写姓名
                if pd.isna(row.姓名):
                    val_list=[]
                    get_str=str1+str2
                    get_str=get_str.replace(str(row.学校),'')
                    get_str=get_str.replace(str(row.班级),'')
                    get_str=get_str.replace(str(row.县区),'')
                    get_str=get_str.replace('昆明','')
                    get_str=get_str.replace('征文','')
                    get_str=get_str.replace('党史','')
                    get_str=get_str.replace('龙区','')
                    get_str=get_str.replace('童心向党','')
                    get_str=get_str.replace('童心永向党','')
                    val_list = val_list+re.findall(name_pattern, get_str)
                    # val_list = val_list+re.findall(name_pattern, str2)

                    if len(val_list)>0:
                        # 使用 filter 函数删除符合条件的元素  使用逻辑取反保留不符合条件的元素
                        condition = lambda x: '从小学' in x or x =='小学' or x =='（小学' or '征文' in x\
                              or '党史“征文' in x or '.' in x or '伟大' in x or '建党' in x  or '精神' in x\
                              or '童心' in x or '时代好队员' in x or '向党' in x or '党话' in x or '温历史' in x
                        val_list = list(filter(lambda x: not condition(x), val_list))
                        #删除列表中已经被其他元素包含了的元素
                        val_list=FieldsInformation.clear_val_list(val_list)
                        # 填写到df中
                        df.at[row.Index, '姓名'] = '|'.join(val_list)
                #填写班级
                # if pd.isna(row.班级):
                #     val_list=[]
                #     val_list =val_list+re.findall(class_pattern, str1)
                #     val_list = val_list+re.findall(class_pattern, str2)
                #     if len(val_list)>0:
                #         #删除列表中已经被其他元素包含了的元素
                #         val_list=FieldsInformation.clear_val_list(val_list)
                #         val_list=list(set(val_list))
                #         # 使用 filter 函数删除符合条件的元素  使用逻辑取反保留不符合条件的元素
                #         condition = lambda x: '从小学' in x or x =='小学' or x =='（小学' or '征文' in x or '党史“征文' in x or '.' in x
                #         val_list = list(filter(lambda x: not condition(x), val_list))
                #         # 填写到df中
                #         df.at[row.Index, '班级'] = '|'.join(val_list)
                #填写组别
                if '初中' in school_info or (pd.notna(row.班级) and any(grade in row.班级 for grade in ['初', '七', '八', '九'])):
                    df.at[row.Index, '组别'] = '初中'
                elif '高中' in school_info or (pd.notna(row.班级) and '高' in row.班级):
                    df.at[row.Index, '组别'] = '高中'
                elif ('小学' in school_info and '从小学' not in school_info) or (pd.notna(row.班级) and any(grade in row.班级 for grade in ['一', '二', '三', '四', '五', '六'])):
                    df.at[row.Index, '组别'] = '小学'
        df.to_excel(out_pathfile, sheet_name=sheet_name, index=True)
    #清洗列表
    def clear_val_list(input_list):
        #删除'\\'
        print(input_list)
        result = [item.replace('\\','') for item in input_list]
        #删除列表中已经被其他元素包含了的元素
        for element in input_list:
            # 检查该元素是否被其他元素包含  
            if not any(element == other and element in other for other in input_list):  
                result.append(element)  # 只有未被包含的元素才会被添加
        
        return result 
#收集文件信息
class CollectFileInformation (): 
    def main(path_name,save_pathfile,para):
        '''
        path：被操作的文件所在的文件夹
        savePathFile：如果要生成文件，如：returnData='excel'时，则该项为被保存的文件路径及文件名
        para：参数
        '''
        #获取 文件路径,文件名,扩展名,打开文件 
        df_result=pd.DataFrame(columns=['文件路径','文件名','扩展名','打开文件'], dtype='object')
        for root, dirs, files in os.walk(path_name, topdown=False):
            for filename in files:
                file_path = os.path.join(root, filename)  
                file_name, file_extension = os.path.splitext(filename)
                df_result.loc[len(df_result)]={
                    '文件路径':root,
                    '文件名':file_name,
                    '扩展名':file_extension,
                    '打开文件':f'=HYPERLINK("{file_path}", "打开文件")' 
                }
        #参数赋值
        file_attributes,narrow_filed,text_analysis,re_spec_dict,\
        specified_before_after,paragraph_text=para
        #加入df_result列名
        df_result = df_result.reindex(columns=
            df_result.columns.tolist() + 
            file_attributes+ #
            narrow_filed+
            list(re_spec_dict.keys())+
            list(specified_before_after.keys())+
            list(paragraph_text.values())
        )
        df_result = df_result.astype('object')  # 或逐列指定类型 df_result['列名'] = df_result['列名'].astype('object')  
        # 获取Doc文档信息（总字数、段落数），同时根据参数统计、分析文档内容
        text_contents=[] #用于存储所有读取的文本内容，以便后续使用
        for index, row in df_result.iterrows():
            print(datetime.now())
            file=os.path.join(row['文件路径'],row['文件名']+row['扩展名'])
            #获取文件属性
            if file_attributes:
                attri=CollectFileInformation.get_file_attributes(file,file_attributes)
                for k,v in attri.items():
                    df_result.loc[index,k]=v
            #获取文本内容，如：doc、docx、txt、pdf
            read_text_content=CollectFileInformation.get_readable_document_content(file)
            text_content=read_text_content['bodyText']+read_text_content['boxText']+read_text_content['table'] #'boxText','pageCount','table'
            text_contents.append(tuple(text_content))#把有文本内容存储下来 未去除首尾空格和空行
            # 去除首尾空格和空行
            text_content=[item.strip() for item in text_content if item.strip() != '']
            #统计、分析文档内容
            CollectFileInformation.statistics_document_content(
                text_content, #已经去除首尾空格和空行的文本内容
                df_result, 
                index, # df_result行号
                narrow_filed,text_analysis,re_spec_dict,specified_before_after,paragraph_text
                )
        #GAI统计、分析文档内容
        # print(df_result.shape, len(text_contents))
        if text_analysis['GAI'] and text_contents:
            for k in list(text_analysis['GAI'].keys()): #k='DeepSeek'、'KiMi'等
                res=TextAnalysis.AI_analysis(text_contents,text_analysis['GAI'][k],k)
                if res:
                    df_result = df_result.assign(**res) # 使用 assign 直接添加新列
                # else:
                #     df_result = df_result.assign(f'AI分析（{AI_name}）'=np.nan)        
            # if 'DeepSeek' in list(text_analysis['GAI'].keys()):
            #     res=TextAnalysis.AI_analysis(text_contents,text_analysis['GAI']['DeepSeek'],'DeepSeek')
            #     df_result = df_result.assign(**res) # 使用 assign 直接添加新列
            # if 'KiMi' in list(text_analysis['GAI'].keys()):
            #     res=TextAnalysis.AI_analysis(text_contents,text_analysis['GAI']['KiMi'],'KiMi')
            #     df_result = df_result.assign(**res) # 使用 assign 直接添加新列
        #生成Excel表
        df_result.to_excel(save_pathfile, index=False)  # index=False 表示不写入行索引  
        print('文件信息收集表保存在：',save_pathfile)
    #供线程使用的函数
    def process_readable_document(get_list): #file_path,kz_name
        file_path=get_list[0]
        suffix_name=get_list[1]
        result=get_list[2]
        # result={'error':[],'bodyText':[],'boxText':[],'pageCount':[],'table':[],}
        if suffix_name=='.docx' or suffix_name=='.docx':
            result = Opration_Word.readWord(file_path)
        elif suffix_name=='.pdf':
            result['bodyText']=Opration_PDF.read_pdf(file_path)
        elif suffix_name=='.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:  # 打开 TXT 文件  
                    text = file.read()  # 读取整个文件的内容  
                    result['bodyText']=text.split('\n')
            except Exception as e:
                print(f"读取 TXT 时出错: {e}")  
                return None  # 如果出错，则返回 None  
        return result
    #获取文件属性（'最后一次修改时间','大小（KB）'等等）
    def get_file_attributes(file_path,file_attributes):
        '''
        filePath：含路径的文件名称
        file_attributes:['大小（KB）','最后修改时间','创建时间','最后访问时间','文件权限']
        '''
        attr={}
        if '大小（KB）' in file_attributes:
            attr['大小（KB）']=round(os.path.getsize(file_path)/1024,2) #'%.2fKB' %(statinfo.st_size/1024)
        if '最后修改时间' in file_attributes:
            attr['最后修改时间']=time.strftime("%Y-%m-%d %H:%M:%S",time.localtime(os.path.getmtime(file_path)))
        if '创建时间' in file_attributes:
            attr['创建时间']=time.strftime("%Y-%m-%d %H:%M:%S",time.localtime(os.path.getctime(file_path)))
        if '最后访问时间' in file_attributes:
            attr['最后访问时间']=time.strftime("%Y-%m-%d %H:%M:%S",time.localtime(os.path.getatime(file_path)))
        if '文件权限' in file_attributes:
            attr['文件权限']=oct(os.stat(file_path).st_mode)[-3:]
        return attr
    #获取可读文档的文本内容，如：doc、docx、txt、pdf
    def get_readable_document_content(file_path):
        result_read={'error':[],'bodyText':[],'boxText':[],'pageCount':[],'table':[],}
        suffix_name=os.path.splitext(file_path)[1].lower()
        if suffix_name=='.doc' or suffix_name=='.docx' or suffix_name=='.pdf' or suffix_name=='.txt':
            # 启动线程执行
            with ThreadPoolExecutor() as executor:
                future = executor.submit(CollectFileInformation.process_readable_document, [file_path,suffix_name,result_read])
                result_read = future.result() # 获取结果，字典
        return result_read #字典
    #分析、统计文档内容
    def statistics_document_content(
            text_contents, #已经去除首尾空格和空行的文本内容
            df_result,
            df_result_index,
            narrow_filed, #常规字段
            text_analysis, #文本分析
            re_spec_dict, # 采用正则表达是的特殊字符
            specified_before_after, # 指定字符前后
            paragraph_text,#段落内容的行号
            ):
        #常规字段，如：'总字数（参考）','有文字的段落数（参考）'
        if narrow_filed:
            df_result.loc[df_result_index,'总字数（参考）']=len(''.join(text_contents))
            df_result.loc[df_result_index,'有文字的段落数（参考）']=len(text_contents)
        #文本分析
        if text_analysis['SnowNLP'] and text_contents:
            res=TextAnalysis.SnowNLP_analysis(''.join(text_contents),text_analysis['SnowNLP'])
            for k,v in res.items():
                df_result.loc[df_result_index,k]=' '.join(v) if type(v)==list else v
        #需要分段统计的
        for index,txt in enumerate(text_contents):
            # 采用正则表达是的特殊字符
            if re_spec_dict:
                for k,v in re_spec_dict.items():
                    x=re.findall(v, txt)
                    if x:
                        df_result.loc[df_result_index,k]=''.join(x[0])
            # 指定字符及其之 前 的指定数量的字符
            if specified_before_after:
                for k,v in specified_before_after.items():
                    if k in txt:
                        if v[0]=='before':
                            if v[2]>0: #指定数量字符
                                df_result.loc[df_result_index,k]=txt[txt.index(k)-v:txt.index(k)+len(k)]
                            elif v[2]==-1: #整段
                                df_result.loc[df_result_index,k]=txt[:txt.index(k)+len(k)]
                        if v[0]=='after':
                            if v[2]>0: #指定数量字符
                                df_result.loc[df_result_index,k]=txt[txt.index(k):txt.index(k)+v]
                            elif v[2]==-1: #整段
                                df_result.loc[df_result_index,k]=txt[txt.index(k):]

            # 段落内容
            if paragraph_text:
                if index+1 in list(paragraph_text.keys()):
                    df_result.loc[df_result_index,paragraph_text[index+1]]=txt

#从申报表中收集信息
class GetInformationFromTable():
    def main(df1,df2,save_file_path):
        '''
        从申报表中采集信息
        '''
        for index,row in df1.iterrows():
            if pd.notna(row.姓名):
                index_list = df2[df2['姓名'] == row.姓名].index
                if not index_list.empty and pd.notna(index_list[0]):
                    if pd.isna(df1.at[index,'学校']):
                        df1.at[index,'学校']=df2.at[index_list[0],'学校和班级']
                    if pd.isna(df1.at[index,'班级']):
                        df1.at[index,'班级']=df2.at[index_list[0],'学校和班级']
                    if pd.isna(df1.at[index,'指导教师']):
                        df1.at[index,'指导教师']=df2.at[index_list[0],'指导教师']
                        df1.at[index,'联系电话']=df2.at[index_list[0],'联系电话']
                    if pd.isna(df1.at[index,'征文标题']):
                        df1.at[index,'征文标题']=df2.at[index_list[0],'征文标题']
                else:
                    pass
                    print(row.姓名)
        df1.to_excel(save_file_path, index=True)
#筛选重复记录
class FilterDuplicateRecords():
    def main(file1,sheet_name,filter_field,save_duplicates,save_unique):
        '''
        筛选记录，如：疑似重复的数据
        '''
        df1=pd.read_excel(file1,sheet_name=sheet_name, header=0, index_col=0)
        # 筛选出相同的记录  
        duplicates = df1.groupby(filter_field).filter(lambda x: len(x) > 1)  
        # 找出重复数据，保留第一条记录
        df_unique = df1.drop_duplicates(subset=filter_field, keep='first')
        #保存结果
        if duplicates.shape[0]>0:
            duplicates.to_excel(save_duplicates, index=True) #重复的所有记录
            df_unique.to_excel(save_unique, index=True) #清除重复行（仅保留第一行）后的所有数据
            print('重复的所有记录保存在：%s；\n清除重复行（仅保留第一行）后的所有数据：%s' 
                  %(save_duplicates,save_unique))
        else:
            print('没有重复的行！(%s)' %('、'.join(filter_field)))
#获得文本的相似度
class GetTextSimilarity():
    def main(file1,sheet_name,text_ID,file_info,save_file):
        import os
        import pandas as pd
        from concurrent.futures import ThreadPoolExecutor
        begin_time=datetime.now()
        df=pd.read_excel(file1,sheet_name=sheet_name, header=0)
        for i in file_info:
            if i not in df.columns:
                print('Excel表的表头必须包含“%s”' %('、'.join(file_info)))
                return
        #将word文档内容保存到字典中
        doc_contents=GetTextSimilarity.get_word_contents(df)
        print('word文档内容读入内存用时：',datetime.now()-begin_time)
        #计算相似度
        reslut_simil=[0]*df.shape[0] #相似度
        reslut_simil_link=['']*df.shape[0] #相似文档链接
        reslut_simil_numb=['']*df.shape[0] #相似文档编号
        reslut_simil_at=['']*df.shape[0] #相似文档位置
        for index1,val1 in doc_contents.items():
            for index2,val2 in doc_contents.items():
                str1=val1[1]
                str2=val2[1]
                if index2!=index1 and str1!='' and str2!='':
                    resl_simi=GetTextSimilarity.txets_Jaccard(str1,str2)
                    if resl_simi>reslut_simil[index1]:
                        reslut_simil[index1]=resl_simi
                        reslut_simil_link[index1]=f'=HYPERLINK("{val2[0]}", "打开相似文件")'
                        reslut_simil_numb[index1]=str(df.loc[index2,text_ID]).zfill(4) #相似文档编号
                        reslut_simil_at[index1]=val2[0] #相似文档位置
        df['相似度']=[str(round(i*100,1))+'%' if i!=0 else'' for i in reslut_simil]
        df['打开相似文件']=reslut_simil_link
        df['相似文档编号']=reslut_simil_numb
        df['相似文档位置']=reslut_simil_at
        #重写打开文件的列
        if '打开文件' in df.columns:
            df['打开文件'] = df.apply(lambda row: f'=HYPERLINK("{row["文件路径"]}\\{row["文件名"]}{row["扩展名"]}", "打开文件")', axis=1)  
        df.to_excel(save_file,index=False)
        print('相似度检查结果，保存在：%s' %(save_file))
        print('全部完成用时：',datetime.now()-begin_time)
    #把word文档的内容保存到字典
    def process_doc_file(file_path):
        import os  
        from win32com import client
        import tempfile    
        from docx import Document
        import pythoncom
        import psutil
        docx_file=file_path
        doc_contents = {'error': [], 'bodyText': [], 'boxText': [], 'table': []}  
        extension = os.path.splitext(file_path)[1].lower()
        #doc文件转换为docx文件
        if extension == '.doc':
            temp_dir = tempfile.gettempdir()
            docx_file = os.path.join(temp_dir, os.path.basename(file_path).replace('.doc', '.docx'))
            docx_file=Opration_Word.doc2docx(file_path,docx_file)
        # 处理 .docx 文件  
        docx = Document(docx_file)
        for paragraph in docx.paragraphs:  
            text = paragraph.text.strip()  
            if text:
                doc_contents['bodyText'].append(text)  
        # 提取文本框中的文本  
        for child in docx.element.body.iter():
            if child.tag.endswith('textbox'):
                text_box_content = ''.join(  
                    c.text for c in child.iter() if c.tag.endswith('main}r')  
                ).strip()
                if text_box_content:  
                    doc_contents['boxText'].append(text_box_content)  
        # 获取表格内容
        for table in docx.tables:  
            for row in table.rows:  
                for cell in row.cells:  
                    text = cell.text.strip()  
                    if text:
                        doc_contents['table'].append(text)
        if extension == '.doc' and os.path.exists(docx_file):  
                os.remove(docx_file)  # 删除文件  
        return file_path, doc_contents
    def get_word_contents(df):
        from concurrent.futures import ThreadPoolExecutor, as_completed  
        result_doc_contents = {}  
        with ThreadPoolExecutor() as executor: #创建线程池
            futures = {} #用于存储提交给线程池的任务和它们对应的索引
            for index1, row1 in df.iterrows():
                # print(index1,row1.扩展名)
                if 'DOC' in str(row1.扩展名).upper():
                    file_path = os.path.join(str(row1.文件路径), str(row1.文件名) + str(row1.扩展名))  
                    if os.path.isfile(file_path):
                        #提交一个任务到线程池，调用 GetTextSimilarity.process_doc_file 函数处理该文件，并将返回的 Future 对象与当前行的索引 index1 关联
                        futures[executor.submit(GetTextSimilarity.process_doc_file, file_path)] = index1  
            print('文件内容读取完毕！')
            # 收集结果
            for future in as_completed(futures):
                index1 = futures[future]
                # print(index1)
                try:
                    file_path, doc_contents = future.result()
                    # print(index1,file_path)
                    result_doc_contents[index1] = [  
                        file_path,
                        ''.join(doc_contents['bodyText']) +  
                        ''.join(doc_contents['table']) +  
                        ''.join(doc_contents['boxText'])  
                    ]
                except Exception as e:
                    print(e)
                    result_doc_contents[index1] = [file_path, f'Error processing file: {str(e)}']  
            print('文件内容收集完毕！')
        return result_doc_contents
    #莱文斯坦距离(适合短文本比较，给出直接的字符串编辑需求)
    def txets_Levenshtein(str1,str2):
        '''
        适合短文本比较，给出直接的字符串编辑需求。
        莱文斯坦距离（Levenshtein distance）是用于衡量两个字符串之间的差异的指标，
        它表示将一个字符串转换为另一个字符串所需的最少单字符编辑操作的数量，
        操作包括插入、删除或替换字符。莱文斯坦距离越大，表示差异越大。
        
        0：完全相同。
        距离小于 3（对于较短字符串，通常 5 个字符以内）：很高的相似度。
        距离在 3 到 5 之间：一个适度的相似度，文字可能有少量错误。
        距离大于 5：相似度较低，字符串差异明显。
        '''
        import Levenshtein  
        distance = Levenshtein.distance(str1, str2)  
        similarity = 1 - distance / max(len(str1), len(str2))  
        # print(f"莱文斯坦距离: {distance}")  
        # print(f"相似度: {similarity:.2f}")
        return similarity
    #余弦相似度(可以捕捉文本之间的语义关系，更适合较长且内容丰富的文本比较)
    def txets_cosine(str1,str2):
        '''
        可以捕捉文本之间的语义关系，更适合较长且内容丰富的文本比较。
        大于 0.9：非常高的相似度。两者几乎完全相同，通常用于文本内容非常接近的情况。
        0.7 到 0.9：高相似度。向量之间有强烈的相似迹象，适用于相似文档或产品推荐。
        0.5 到 0.7：中等相似度。有一定的相似性，但不算特别接近。
        0 到 0.5：低相似度。向量之间几乎没有相似性，可能表示不同的主题或内容。
        '''
        from sklearn.feature_extraction.text import CountVectorizer  
        from sklearn.metrics.pairwise import cosine_similarity    
        vectorizer = CountVectorizer().fit_transform([str1,str2])  
        vectors = vectorizer.toarray()
        return cosine_similarity(vectors)[0][1]
    # Jaccard相似度(通过计算文本中的独立词汇，适合比较内容大致相似的长文本。)
    def txets_Jaccard(str1,str2):
        '''
        通过计算文本中的独立词汇，适合比较内容大致相似的长文本。
        计算的是两个集合交集和并集的比值，常用于文本比较。可通过计算字符串中词的集合来应用。
        值范围在0到1之间
        小于 0.2：相似度低，集合几乎没有交集。
        0.2 到 0.4：相似度一般，有一定的重叠，但不算高。
        0.4 到 0.6：相似度较高，两个集合有显著的共同元素。
        0.6 到 0.8：相似度很高，集合间有许多共同元素。
        大于 0.8：非常高的相似度，集合几乎相同。
        '''
        set1 = set(str1)
        set2 = set(str2)
        intersection = len(set1.intersection(set2)) #找到两个集合（或数据结构）之间的交集
        union = len(set1.union(set2)) #获取两个集合的并集
        # print(f"Jaccard相似度: {intersection / union}")
        return intersection / union

#用读取Excel表中的文件所在路径及文件名，在新的文件夹生成新文件名的文件
class GenerateNewFilename(): 
    def main(excelPathFile,newPath,newNameCombo,sheet_name):
        df=pd.read_excel(excelPathFile,sheet_name=sheet_name, header=0)
        df['编号'] =  df['编号'].astype(str).str.zfill(4) 
        if '扩展名' not in df.columns or '文件路径' not in df.columns or '文件名' not in df.columns:
            print('Excel文件指定sheet中须包含“路径、文件名、后缀”')
            return
        for i in newNameCombo[1]:
            if i not in df.columns:
                print('Excel文件指定sheet中须包含“%s”' %(i))
                return
        #拷贝文件
        for index,row in df.iterrows():
            oPathName=os.path.join(row.文件路径,row.文件名+row.扩展名)
            if os.path.exists(oPathName)==False:
                print('未找到文件：%s' %(oPathName),str(row.编号))
                return
            if os.path.exists(newPath)==False:
                os.makedirs(newPath)
            newName=[]
            for i in newNameCombo[1]:
                newName.append(row[i])
            newName=newNameCombo[0].join(newName)+row.扩展名
            shutil.copy(oPathName, os.path.join(newPath,newName))
        # 检验将拷贝后的文件名进行回收，以便检验
        df1 = pd.DataFrame(columns=newNameCombo[1])
        for root, dirs, files in os.walk(newPath):  
            for file in files:  
                # 打印文件的完整路径
                x=file.split('_')
                df1.loc[len(df1)]=x #注意文件名里自身就有的下划线
        df1.to_excel('E:\\1\\总表.xlsx',index=False)
        # =HYPERLINK("\征文\"&A2&"_"&B2&"_"&C2, "打开文件")
#批量发送电子邮件
class SendEmail():
    def SendEmail1():
        import zmail,os
        # 发件人（用户名、密码）
        sender =('69301020@qq.com','dceyjpkswzhhbhbf')
        path=r'E:\20230109中华魂演讲、征文作品\20240117征文（送审）\征文（送审）\分解并发送电子邮件给评委'
        allDict={
            # '张向东':[('look.east@163.com'),('1征文（送审）.rar',),()],
            # '张建欣':[('kmswgy33@126.com'),('2征文（送审）.rar',),()],
            # '杨丽': [('653285204@qq.com'),('3征文（送审）.rar',),()],
            # '张玉芬': [('857167452@qq.com'),('4征文（送审）.rar',),()],
            # '陈华':[('154596294@qq.com'),('5征文（送审）.rar',),()],
            # '徐九林':[('69301020@qq.com'),('6征文（送审）.rar',),()],
        }

        for i in list(allDict.keys()):
            subject="“中华魂”（毛泽东伟大精神品格）征文评选"
            content_text="尊敬的%s老师，您好！\n \
                “中华魂”（毛泽东伟大精神品格）征文活动已启动。现征文评选工作面临时间紧，任务重，望各位评委老师谅解、大力支持！\n \
                评分前请认真阅读“评分细则”，并将评分结果填写到 “中华魂”（毛泽东伟大精神品格）征文初评表 中。" \
                %(i)
            revicer = allDict[i][0] # 收件人
            attachments=[] # 附件,多个附件，以列表的形式存储
            for i1 in allDict[i][1]:
                attachments.append(os.path.join(path,i1))
            copy = list(allDict[i][2]) # 抄送人
            print(copy)
            # 邮件内容 - 必须以字典来存储
            mail_msg = {
                "subject":subject,# 主题
                "content_text":content_text,# 正文-纯文本
                'attachments':attachments # 附件,多个附件，以列表的形式存储
            }
            # 发送邮件
            # 1-构建发送邮件的服务
            # server = zmail.server(sender[0],sender[1])
            server = zmail.server(*sender)
            # 2-通过邮件的服务去发送邮件
            server.send_mail(recipients=revicer,
                            mail=mail_msg,
                            cc=copy)
#按照Excel指定的两列（序号、分类文件夹名 如：获奖等第）将当前文件分类到指定文件夹中
class DirClassify():
    def dirClassify1(excelPathFile,fileNameCol,classifyCol,filePath):
        #判断文件及文件夹是否存在
        if os.path.exists(excelPathFile)==False:
            print("Excel文件不存在！")
            return
        if os.path.exists(filePath)==False:
            print("保存文件的文件夹不存在！")
            return
        #获取文件夹中所有文件
        allFilenames=os.listdir(filePath)
        print(len(allFilenames))
        #读取Excel表
        xlsx=ReadExcel_pandas(excelPathFile)
        #表头在第一张sheet的第2行
        xlsxHeader=xlsx.values[0][0]
        xlsxValues=xlsx.values[0][1:]
        if fileNameCol not in xlsxHeader:
            print('指定的列名出错！')
            return
        for i in classifyCol:
            if i not in xlsxHeader:
                print('指定的列名出错！')
                return
        fileNameNum=xlsxHeader.index(fileNameCol)
        for i in xlsxValues:
            if i[fileNameNum]:
                for i1 in classifyCol:
                    classifyColNum=xlsxHeader.index(i1)
                    if i[classifyColNum]!='':
                        for i2 in allFilenames:
                            if i2.startswith(i[fileNameNum]):
                                if os.path.isdir(filePath+'//'+i[classifyColNum])==False:
                                    os.mkdir(filePath+'//'+i[classifyColNum])
                                oPathName=os.path.join(filePath,i2)
                                if os.path.exists(oPathName)==False:
                                    print("文件不存在（%s）"%(oPathName))
                                    continue
                                nPathName=os.path.join(filePath+'//'+i[classifyColNum],i2)
                                shutil.copy(oPathName, nPathName)


Main.main()
# if __name__ == "__main__":
#     root = tk.Tk()
#     app = CollectionAndAnalysisOfFileInformation(root)
#     root.mainloop()
